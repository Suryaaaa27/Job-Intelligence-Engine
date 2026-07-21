"""Extracts raw text from DOCX resumes."""

try:
    import docx
except ModuleNotFoundError:  # pragma: no cover - optional dependency
    docx = None


def _iter_paragraphs_and_tables(document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            for para in document.paragraphs:
                if para._p is child:
                    if para.text.strip():
                        yield para.text.strip()
                    break
        elif child.tag.endswith("}tbl"):
            for table in document.tables:
                if table._tbl is child:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                yield cell.text.strip()
                    break


def extract_text_from_docx(file_path: str) -> str:
    if docx is None:
        raise RuntimeError("python-docx is required to parse DOCX resumes. Install resume_parser/requirements.txt")
    document = docx.Document(file_path)
    lines = list(_iter_paragraphs_and_tables(document))
    return "\n".join(lines)


def extract_lines_from_docx(file_path: str) -> list:
    if docx is None:
        raise RuntimeError("python-docx is required to parse DOCX resumes. Install resume_parser/requirements.txt")
    document = docx.Document(file_path)
    return list(_iter_paragraphs_and_tables(document))
